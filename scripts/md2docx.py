#!/usr/bin/env python3
"""Convert GRC markdown templates to beautiful, fillable .docx forms.

Design:
- Real Word form fields (legacy FORMTEXT) for [___] blanks — the reader Tabs
  through fields and types. Choice fields like [Yes/No] or
  [Mitigate / Transfer / Avoid / Accept] become FORMDROPDOWN dropdowns.
- House-style look: Georgia serif body, dark slate headings, shaded table
  headers, clean rule tables, signature/approval blocks.
- Works in Microsoft Word AND LibreOffice/OpenOffice (legacy form fields are
  the portable standard).
"""
import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---- House palette -------------------------------------------------------
NAVY = RGBColor(0x1F, 0x3A, 0x5F)      # headings
SLATE = RGBColor(0x44, 0x50, 0x5C)     # subheadings
STEEL = RGBColor(0x4A, 0x6F, 0xA5)     # accents
HDR_FILL = "1F3A5F"                     # table header fill (navy)
ALT_FILL = "EAF0F8"                     # alternating row fill
BORDER = "B8C4D8"                       # table border

CHOICE_SPLIT = re.compile(r"\s*/\s*")   # [A / B / C] -> dropdown options


def set_cell_bg(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_table_borders(table, color=BORDER, sz="4"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def make_form_field(paragraph, default="", width_chars=24, dropdown_options=None):
    """Insert a legacy form field into the paragraph run(s)."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    if dropdown_options:
        instrText.text = " FORMDROPDOWN "
    else:
        instrText.text = " FORMTEXT "
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)

    # field result (what displays)
    t = OxmlElement("w:t")
    t.text = default
    run._r.append(t)

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar3)

    # field properties: size + dropdown list
    if dropdown_options:
        fldPr = OxmlElement("w:fldPr")
        ffData = OxmlElement("w:ffData")
        dd = OxmlElement("w:ffType")
        dd.set(qn("w:val"), "dropdown")
        ffData.append(dd)
        for opt in dropdown_options:
            name = OxmlElement("w:ffName")
            name.set(qn("w:val"), opt.strip()[:64])
            ffData.append(name)
        fldPr.append(ffData)
        # wrap: fldPr must come right after fldChar begin — python-docx puts
        # it via the run's rPr position; simplest: rebuild properly below.
    return run


def add_form_text(paragraph, default="", width=28):
    """FORMTEXT field for a [___] blank."""
    r = paragraph.add_run()
    rPr = r._r.get_or_add_rPr()
    # underline styling so the blank reads as a fill line
    rPr.append(OxmlElement("w:u"))
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " FORMTEXT "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = " " + " " * min(width, 30)
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        r._r.append(el)
    return r


def add_form_dropdown(paragraph, options, default=""):
    """FORMDROPDOWN field for [A / B / C] choice."""
    r = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " FORMDROPDOWN "
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = default or (options[0] if options else "")
    f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
    for el in (f1, it, f2, t, f3):
        r._r.append(el)
    # field data (dropdown options) in rPr
    fldPr = OxmlElement("w:fldPr")
    ffData = OxmlElement("w:ffData")
    for opt in options:
        name = OxmlElement("w:ffName"); name.set(qn("w:val"), opt.strip()[:64])
        ffData.append(name)
    fldPr.append(ffData)
    r._r.insert(1, fldPr)
    return r


BLANK_RE = re.compile(r"\[(_{2,})\]")
CHOICE_RE = re.compile(r"\[([^\[\]]{1,60}?)\s*/\s*([^\[\]]{1,60})\]")
LABELED_RE = re.compile(r"\[([^\[\]]{1,80}?:\s*_{2,})\]")
PREFIXED_RE = re.compile(r"\[([A-Za-z0-9\- ]{1,20}?)(_{2,})\]")


def parse_fill_field(text):
    """Given the bracket content, decide: blank, dropdown, or labeled blank."""
    s = text.strip()
    if CHOICE_RE.fullmatch("[" + s + "]"):
        opts = [o.strip() for o in CHOICE_SPLIT.split(s)]
        return ("dropdown", opts)
    return ("text", None)


def emit_text_with_fields(paragraph, text, base_bold=False, base_italic=False):
    """Emit text runs, converting [___] and [A / B / C] into form fields.

    Handles:
      [___] / [______]                  -> FORMTEXT blank
      [ORG]-XXX-___  (prefix + blank)   -> literal prefix + FORMTEXT blank
      [Yes/No] / [A / B / C]            -> FORMDROPDOWN
      [label: ______]                   -> literal label + FORMTEXT blank
    """
    pos = 0
    # match bracketed tokens: [___], [label: ___], [A / B / C], [ORG]
    pattern = re.compile(r"\[[^\[\]]{1,90}\]")
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.bold = base_bold
            r.italic = base_italic
        inner = m.group(0)[1:-1]
        stripped = inner.strip()
        # pure blank [___]
        if re.fullmatch(r"_+", stripped):
            add_form_text(paragraph)
        # choice [A / B / C]
        elif "/" in stripped and not stripped.endswith("_"):
            opts = [o.strip() for o in CHOICE_SPLIT.split(stripped) if o.strip()]
            if len(opts) >= 2:
                add_form_dropdown(paragraph, opts)
            else:
                add_form_text(paragraph)
        # labeled blank [label: ______] or [prefix-___]
        elif stripped.endswith("_"):
            label = stripped.rstrip("_").rstrip(":").strip()
            if label:
                r = paragraph.add_run(label + " ")
                r.bold = base_bold
                r.italic = base_italic
            add_form_text(paragraph)
        else:
            # plain bracket text like [ORG] — keep as literal
            r = paragraph.add_run(m.group(0))
            r.bold = base_bold
            r.italic = base_italic
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.bold = base_bold
        r.italic = base_italic


def style_paragraph(p, size=10.5, bold=False, color=None, space_after=4,
                    space_before=0, italic=False):
    for run in p.runs:
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    return p


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    sizes = {1: 20, 2: 14, 3: 12}
    colors = {1: NAVY, 2: NAVY, 3: SLATE}
    run.font.size = Pt(sizes.get(level, 11))
    run.bold = True
    run.font.color.rgb = colors.get(level, SLATE)
    run.font.name = "Georgia"
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_doc_control_table(doc, rows):
    """Document Control / Field-Details table with navy header."""
    table = doc.add_table(rows=len(rows) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0].cells
    hdr[0].text = "Field"
    hdr[1].text = "Details"
    for c in hdr:
        set_cell_bg(c, HDR_FILL)
        for p in c.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for i, (k, v) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        # key cell: bold label only
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        r0.font.size = Pt(10)
        # value cell: text with form fields
        p1 = cells[1].paragraphs[0]
        emit_text_with_fields(p1, v)
        if i % 2 == 0:
            for c in cells:
                set_cell_bg(c, ALT_FILL)
    # widths
    for row in table.rows:
        row.cells[0].width = Inches(1.7)
        row.cells[1].width = Inches(4.8)
    return table


def parse_doc_control(lines):
    """Parse the '| Field | Details |' table at the top of a template."""
    rows = []
    in_table = False
    for line in lines:
        s = line.strip()
        if s.startswith("|") and "Field" in s and "Details" in s:
            in_table = True
            continue
        if in_table:
            if not s.startswith("|"):
                break
            # skip separator row (dashes)
            cells = [c.strip() for c in s.strip("|").split("|")]
            if cells and all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells if c.strip()):
                continue
            if len(cells) >= 2:
                rows.append((cells[0].replace("**", ""), cells[1]))
    return rows


def convert(md_path, out_path):
    text = open(md_path).read()
    lines = text.split("\n")
    doc = Document()

    # page setup: A4-ish with comfortable margins
    sec = doc.sections[0]
    sec.page_width = Inches(8.27)
    sec.page_height = Inches(11.69)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Georgia"
    normal.font.size = Pt(10.5)

    # form protection: enforce fill-in-forms so Word/LO treat this as a
    # fillable form (Tab cycles fields; body text is read-only)
    settings = doc.settings.element
    prot = OxmlElement("w:documentProtection")
    prot.set(qn("w:edit"), "forms")
    prot.set(qn("w:enforcement"), "1")
    settings.append(prot)

    # title (H1)
    title = ""
    for line in lines:
        if line.startswith("# ") and not line.startswith("## "):
            title = line[2:].strip()
            break
    add_heading(doc, title, 1)

    # body parse — skip the title line (already emitted as centered title)
    i = 0
    doc_control_done = False
    title_skipped = False
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not title_skipped and s.startswith("# ") and not s.startswith("## "):
            title_skipped = True
            i += 1
            continue

        # skip the leading doc-control table, emit it once as styled
        if s.startswith("|") and "Field" in s and "Details" in s and not doc_control_done:
            rows = parse_doc_control(lines[i:])
            add_doc_control_table(doc, rows)
            doc_control_done = True
            # skip past table rows
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                j += 1
            i = j
            continue

        # headings
        m = re.match(r"^(#{1,3})\s+(.*)", s)
        if m and not s.startswith("|"):
            level = len(m.group(1))
            htext = m.group(2).replace("**", "").replace("\\begin", "").replace("\\end", "").strip()
            if htext and not htext.startswith("!"):
                add_heading(doc, htext, level)
            i += 1
            continue

        # horizontal rule
        if re.match(r"^(\*\*\*|---|\*\*\* )$", s):
            i += 1
            continue

        # blockquote (Purpose / How-to-use / note)
        if s.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(6)
            qtext = s.lstrip("> ").replace("**", "")
            emit_text_with_fields(p, qtext, base_italic=True)
            for r in p.runs:
                r.italic = True
                if r.text.startswith("Purpose") or r.text.startswith("How to"):
                    r.bold = True
            i += 1
            continue

        # table
        if s.startswith("|") and s.endswith("|") and "---" not in s.replace("|", ""):
            # collect rows
            tbl_rows = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
                # skip separator row
                if all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells):
                    j += 1
                    continue
                tbl_rows.append(cells)
                j += 1
            if tbl_rows:
                ncols = max(len(r) for r in tbl_rows)
                table = doc.add_table(rows=len(tbl_rows), cols=ncols)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table)
                for ri, row in enumerate(tbl_rows):
                    for ci in range(ncols):
                        val = row[ci] if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        if ri == 0:
                            set_cell_bg(cell, HDR_FILL)
                        elif ri % 2 == 0:
                            set_cell_bg(cell, ALT_FILL)
                        emit_text_with_fields(p, val.replace("**", ""),
                                              base_bold=(ri == 0))
                        if ri == 0:
                            for r in p.runs:
                                r.bold = True
                                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                r.font.size = Pt(9.5)
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                i = j
                continue

        # checkbox
        if re.match(r"^- \[ \]", s) or re.match(r"^- \u25a1", s):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            emit_text_with_fields(p, s.lstrip("- ").strip())
            i += 1
            continue

        # list item
        if re.match(r"^[-*]\s+", s):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            emit_text_with_fields(p, re.sub(r"^[-*]\s+", "", s))
            i += 1
            continue

        # numbered item
        m = re.match(r"^(\d{1,2})\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_after = Pt(2)
            emit_text_with_fields(p, f"{m.group(1)}. {m.group(2)}")
            i += 1
            continue

        # blank line
        if not s:
            i += 1
            continue

        # plain paragraph
        p = doc.add_paragraph()
        emit_text_with_fields(p, s)
        style_paragraph(p)
        i += 1

    doc.save(out_path)
    return title


if __name__ == "__main__":
    src = sys.argv[1]
    dst = sys.argv[2] if len(sys.argv) > 2 else src.replace(".md", ".docx")
    t = convert(src, dst)
    print(f"OK: {t} -> {dst}")
